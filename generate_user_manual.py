"""
Script tạo tài liệu hướng dẫn sử dụng ứng dụng Quản lý Kho Pháo Hoa
Output: huong_dan_su_dung.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows, header_color="1A73E8", col_widths=None):
    """Create a formatted table with colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FC")

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point. Optionally bold a prefix portion."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        rest = text[len(bold_prefix):]
        run2 = p.add_run(rest)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_numbered(doc, text, bold_prefix=None):
    """Add a numbered list item."""
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        rest = text[len(bold_prefix):]
        run2 = p.add_run(rest)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_note_box(doc, text, box_type="info"):
    """Add a colored note/tip/warning box."""
    colors = {
        "info":    ("E8F0FE", "1A73E8", "ℹ️  Lưu ý"),
        "tip":     ("E6F4EA", "1E8E3E", "💡 Mẹo"),
        "warning": ("FFF3E0", "E65100", "⚠️  Cảnh báo"),
        "danger":  ("FCE8E6", "C5221F", "🚫 Quan trọng"),
    }
    bg_color, text_color, title = colors.get(box_type, colors["info"])

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)

    p = cell.paragraphs[0]
    title_run = p.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.size = Pt(10)
    r, g, b = int(text_color[0:2], 16), int(text_color[2:4], 16), int(text_color[4:6], 16)
    title_run.font.color.rgb = RGBColor(r, g, b)

    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Add some spacing
    doc.add_paragraph()


def add_body(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


# ─── Main Document ────────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Custom styles ──
    styles = doc.styles

    # Heading 1 style
    h1 = styles["Heading 1"]
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2 style
    h2 = styles["Heading 2"]
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0x19, 0x67, 0xD2)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3 style
    h3 = styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════

    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("HƯỚNG DẪN SỬ DỤNG")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("ỨNG DỤNG QUẢN LÝ KHO PHÁO HOA")
    sub_run.bold = True
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = RGBColor(0x42, 0x85, 0xF4)

    doc.add_paragraph()

    # Decorative line
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run("━" * 40)
    line_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    line_run.font.size = Pt(14)

    doc.add_paragraph()

    # Version and date
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("Phiên bản: 1.0.0\nNgày cập nhật: 03/04/2026")
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc_p.add_run("Tài liệu dành cho người dùng cuối (End User)")
    desc_run.font.size = Pt(11)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Page break
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("MỤC LỤC", level=1)

    toc_items = [
        "1.  Giới thiệu",
        "2.  Đăng nhập & Đăng ký tài khoản",
        "3.  Tổng quan giao diện chính",
        "4.  Tab Tồn kho (Dashboard)",
        "5.  Tab Danh mục",
        "    5.1. Quản lý Sản phẩm",
        "    5.2. Quản lý Khách hàng",
        "6.  Tab Nhật ký (Giao dịch)",
        "    6.1. Tạo phiếu Nhập kho",
        "    6.2. Tạo phiếu Xuất kho",
        "    6.3. Lọc lịch sử giao dịch",
        "    6.4. Báo cáo hàng ngày",
        "7.  Quản lý Công nợ Khách hàng",
        "8.  Đối soát Tồn kho",
        "9.  Thống kê & Báo cáo",
        "10. Cài đặt",
        "11. Phân loại sản phẩm theo Nghị định 137",
        "12. Câu hỏi thường gặp (FAQ)",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        if not item.startswith("    "):
            run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1: GIỚI THIỆU
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("1. Giới thiệu", level=1)

    add_body(doc,
        "Quản lý Kho Pháo Hoa là ứng dụng di động giúp bạn quản lý toàn diện "
        "hoạt động kho hàng pháo hoa. Ứng dụng được thiết kế đơn giản, dễ sử dụng, "
        "phù hợp với nhân viên kho và quản lý."
    )

    add_body(doc, "Các tính năng chính của ứng dụng:")

    features = [
        "Quản lý tồn kho theo 3 kho riêng biệt (Kho 1, Kho 2, Kho 3)",
        "Nhập kho / Xuất kho với theo dõi chi tiết từng sản phẩm",
        "Quản lý danh mục sản phẩm theo phân loại & quy chuẩn pháp lý (NĐ 137)",
        "Quản lý khách hàng (Khách quen / Khách lẻ)",
        "Theo dõi công nợ — Ghi nợ, thanh toán, tất toán",
        "Đối soát tồn kho — Kiểm kê thực tế so với hệ thống",
        "Thống kê & Báo cáo — Biểu đồ nhập/xuất, lợi nhuận, sản phẩm bán chạy",
        "Nhắc nhở hàng ngày — Thông báo tổng kết cuối ngày",
    ]
    for f in features:
        add_bullet(doc, f"✅  {f}")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2: ĐĂNG NHẬP & ĐĂNG KÝ
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("2. Đăng nhập & Đăng ký tài khoản", level=1)

    doc.add_heading("2.1. Đăng nhập", level=2)

    add_numbered(doc, "Mở ứng dụng Quản lý Kho Pháo Hoa")
    add_numbered(doc, "Tại màn hình đăng nhập, nhập Email và Mật khẩu")
    add_numbered(doc, "Nhấn nút Đăng nhập")
    add_numbered(doc, "Nhấn biểu tượng 👁 bên phải ô mật khẩu để hiện/ẩn mật khẩu")

    add_note_box(doc,
        "Nếu nhập đúng thông tin, ứng dụng sẽ tự động chuyển đến màn hình chính. "
        "Nếu sai, sẽ hiện thông báo lỗi phía dưới.", "tip")

    doc.add_heading("2.2. Đăng ký tài khoản mới", level=2)

    add_numbered(doc, 'Tại màn hình đăng nhập, nhấn "Chưa có tài khoản? Đăng ký"')
    add_numbered(doc, "Điền thông tin: Họ và tên (tùy chọn), Email (bắt buộc), Mật khẩu (bắt buộc)")
    add_numbered(doc, "Nhấn nút Đăng ký")
    add_numbered(doc, 'Nếu đã có tài khoản, nhấn "Đã có tài khoản? Đăng nhập" để quay lại')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3: TỔNG QUAN GIAO DIỆN
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("3. Tổng quan giao diện chính", level=1)

    add_body(doc,
        "Sau khi đăng nhập thành công, giao diện chính có 3 tab ở thanh dưới cùng:"
    )

    add_styled_table(doc,
        ["Tab", "Biểu tượng", "Mô tả"],
        [
            ["Tồn kho", "📊 Dashboard", "Xem tổng quan tồn kho từng sản phẩm theo từng kho"],
            ["Danh mục", "📦 Category", "Quản lý sản phẩm & khách hàng"],
            ["Nhật ký", "📖 Journal", "Ghi nhận & xem lại giao dịch nhập/xuất kho"],
        ],
        col_widths=[3, 4, 9],
    )

    add_body(doc,
        "Thanh tiêu đề (phía trên) hiển thị tên ứng dụng và nút ⚙ Cài đặt ở góc phải."
    )

    add_note_box(doc,
        "Tab Nhật ký được mặc định mở trước vì đây là chức năng sử dụng thường xuyên nhất.",
        "info")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4: TAB TỒN KHO
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("4. Tab Tồn kho (Dashboard)", level=1)

    doc.add_heading("4.1. Tổng quan", level=2)

    add_body(doc, "Tại tab Tồn kho, bạn sẽ thấy các thành phần chính sau:")

    doc.add_heading("Thẻ tóm tắt (cuộn ngang phía trên)", level=3)

    add_styled_table(doc,
        ["Thẻ", "Ý nghĩa"],
        [
            ["Tổng sản phẩm", "Số lượng loại sản phẩm đang có trong kho"],
            ["Giá trị kho", "Tổng giá trị hàng tồn (tính theo giá nhập)"],
            ["Sắp hết", "Số sản phẩm có số lượng từ 1–4 (mức nguy hiểm)"],
            ["Hết hàng", "Số sản phẩm có số lượng = 0"],
        ],
        col_widths=[4, 12],
    )

    doc.add_heading("Thanh tìm kiếm", level=3)
    add_body(doc, "Nhập tên sản phẩm để tìm nhanh trong danh sách tồn kho.")

    doc.add_heading("Các tab kho", level=3)
    add_bullet(doc, "Tất cả: Hiển thị tổng tồn kho tất cả các kho", bold_prefix="Tất cả: ")
    add_bullet(doc, "Kho 1 / Kho 2 / Kho 3: Xem tồn kho riêng từng kho", bold_prefix="Kho 1 / Kho 2 / Kho 3: ")

    doc.add_heading("4.2. Xem chi tiết tồn kho sản phẩm", level=2)

    add_numbered(doc, "Nhấn vào bất kỳ sản phẩm nào trong danh sách")
    add_numbered(doc, "Một bảng chi tiết hiện lên, hiển thị tên sản phẩm, số lượng tồn tại từng kho, và lịch sử giao dịch gần đây")

    doc.add_heading("4.3. Trạng thái tồn kho (mã màu)", level=2)

    add_body(doc, "Mỗi sản phẩm được gắn nhãn trạng thái theo số lượng:")

    add_styled_table(doc,
        ["Nhãn", "Màu", "Điều kiện"],
        [
            ["Hết hàng", "🔴 Đỏ", "Số lượng = 0"],
            ["Sắp hết", "🔴 Đỏ", "Số lượng < 5"],
            ["Thấp", "🟡 Vàng", "Số lượng < 20"],
            ["Đủ hàng", "🟢 Xanh", "Số lượng ≥ 20"],
        ],
        header_color="5F6368",
        col_widths=[4, 4, 8],
    )

    doc.add_heading("4.4. Truy cập nhanh", level=2)

    add_bullet(doc, "Nút 📊 Thống kê (góc trên phải): Mở trang thống kê chi tiết", bold_prefix="Nút 📊 Thống kê: ")
    add_bullet(doc, "Nút ✅ Đối soát (nút nổi phía dưới): Mở trang đối soát/kiểm kê tồn kho", bold_prefix="Nút ✅ Đối soát: ")
    add_bullet(doc, "Kéo xuống để làm mới (Pull-to-refresh): Cập nhật dữ liệu mới nhất", bold_prefix="Kéo xuống để làm mới: ")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5: TAB DANH MỤC
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("5. Tab Danh mục", level=1)

    add_body(doc, 'Tab Danh mục được chia thành 2 tab con: "Sản phẩm" và "Khách hàng".')

    doc.add_heading("5.1. Quản lý Sản phẩm", level=2)

    doc.add_heading("Xem danh sách", level=3)
    add_body(doc, "Danh sách hiển thị tất cả sản phẩm với thông tin: Tên sản phẩm, Loại và Phân loại pháp lý (A/B/C), Giá xuất và Giá nhập. Sử dụng thanh tìm kiếm để tìm nhanh theo tên hoặc loại.")

    doc.add_heading("Thêm sản phẩm mới", level=3)
    add_numbered(doc, 'Nhấn nút ➕ (nút tròn phía dưới bên phải)')
    add_numbered(doc, "Điền thông tin: Tên sản phẩm (bắt buộc), Loại, Phân loại pháp lý, Đơn vị tính, Giá nhập, Giá xuất")
    add_numbered(doc, 'Nhấn "Lưu" để hoàn tất')

    add_body(doc, "Các loại sản phẩm hỗ trợ:")
    add_styled_table(doc,
        ["Danh mục", "Mô tả"],
        [
            ["Viên", "Viên (đơn lẻ)"],
            ["Nháy", "Nháy"],
            ["Viên 36", "Viên 36 phát"],
            ["Giàn 25", "Giàn 25 phát"],
            ["Giàn 50", "Giàn 50 phát"],
            ["Giàn 75", "Giàn 75 phát"],
            ["Giàn 100", "Giàn 100 phát"],
            ["Giàn 150", "Giàn 150 phát"],
            ["Giàn 225", "Giàn 225 phát"],
            ["Khác", "Sản phẩm khác"],
        ],
        header_color="34A853",
        col_widths=[4, 12],
    )

    doc.add_heading("Xem & sửa sản phẩm", level=3)
    add_numbered(doc, "Nhấn vào sản phẩm trong danh sách → Mở trang Chi tiết sản phẩm")
    add_numbered(doc, 'Nhấn "Sửa" để chỉnh sửa thông tin')

    doc.add_heading("Xóa sản phẩm", level=3)
    add_bullet(doc, "Cách 1: Vuốt sản phẩm sang trái (swipe) → Sản phẩm sẽ bị xóa", bold_prefix="Cách 1: ")
    add_bullet(doc, "Cách 2: Vào chi tiết sản phẩm → Nhấn nút Xóa", bold_prefix="Cách 2: ")

    add_note_box(doc, "Xóa sản phẩm là không thể hoàn tác. Hãy chắc chắn trước khi thực hiện.", "warning")

    # ── 5.2 Khách hàng ──
    doc.add_heading("5.2. Quản lý Khách hàng", level=2)

    doc.add_heading("Xem danh sách khách hàng", level=3)
    add_body(doc, "Mỗi khách hàng hiển thị: Tên & Số điện thoại, Loại (⭐ Khách quen hoặc 👤 Khách lẻ), Công nợ hiện tại (nếu có), và Trạng thái (✅ xanh nếu không nợ).")

    doc.add_heading("Thêm khách hàng mới", level=3)
    add_numbered(doc, 'Nhấn nút 👤➕ (nút tròn phía dưới bên phải)')
    add_numbered(doc, "Nhập: Tên khách hàng (bắt buộc), Số điện thoại (tùy chọn), Loại (Khách lẻ / Khách quen)")
    add_numbered(doc, 'Nhấn "Thêm"')

    doc.add_heading("Xem chi tiết & quản lý khách hàng", level=3)
    add_body(doc, "Nhấn vào khách hàng → Mở bảng chi tiết với các nút hành động:")

    add_styled_table(doc,
        ["Nút", "Chức năng"],
        [
            ["📋 Xem lịch sử công nợ", "Mở trang chi tiết lịch sử ghi nợ & thanh toán"],
            ["💳 Thanh toán", "Thanh toán một phần công nợ"],
            ["✅ Tất toán", "Xóa toàn bộ công nợ"],
            ["✏️ Sửa", "Chỉnh sửa thông tin khách hàng"],
            ["🗑 Xóa", "Xóa khách hàng (có xác nhận)"],
        ],
        header_color="EA4335",
        col_widths=[5, 11],
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6: TAB NHẬT KÝ
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("6. Tab Nhật ký (Giao dịch)", level=1)

    doc.add_heading("6.1. Tổng quan", level=2)

    add_body(doc,
        "Tab Nhật ký là nơi bạn thực hiện và theo dõi tất cả các giao dịch nhập/xuất kho. "
        "Giao diện bao gồm:")
    add_bullet(doc, "2 nút lớn phía trên: NHẬP KHO (xanh) và XUẤT KHO (cam)")
    add_bullet(doc, "Bộ lọc: Lọc lịch sử giao dịch theo nhiều tiêu chí")
    add_bullet(doc, "Danh sách giao dịch: Liệt kê tất cả giao dịch đã thực hiện")

    doc.add_heading("6.2. Tạo phiếu Nhập kho", level=2)

    add_numbered(doc, "Nhấn nút NHẬP KHO (màu xanh)")
    add_numbered(doc, "Chọn kho: Kho 1, Kho 2, hoặc Kho 3")
    add_numbered(doc, 'Nhập Nguồn nhập / Nhà cung cấp (tùy chọn). VD: "Nhà cung cấp A", "Nhập bổ sung"')
    add_numbered(doc, 'Thêm sản phẩm: Nhấn nút "Thêm" → Chọn sản phẩm → Nhập số lượng')
    add_numbered(doc, "Có thể thêm nhiều sản phẩm trong cùng 1 phiếu. Nhấn ❌ để xóa sản phẩm khỏi phiếu")
    add_numbered(doc, "Xem Tổng cộng giá trị phiếu (tính tự động)")
    add_numbered(doc, "Thêm Ghi chú nếu cần")
    add_numbered(doc, 'Nhấn "Xác nhận nhập kho"')

    add_note_box(doc,
        "Sau khi xác nhận, số lượng tồn kho sẽ tự động được cập nhật tăng tương ứng tại kho đã chọn.",
        "info")

    doc.add_heading("6.3. Tạo phiếu Xuất kho", level=2)

    add_numbered(doc, "Nhấn nút XUẤT KHO (màu cam)")
    add_numbered(doc, "Quy trình tương tự Nhập kho, nhưng:")
    add_bullet(doc, 'Trường "Nơi nhận / Lý do xuất" thay cho "Nguồn nhập"')
    add_bullet(doc, "Giá tính theo giá xuất (thay vì giá nhập)")
    add_numbered(doc, 'Nhấn "Xác nhận xuất kho"')

    add_note_box(doc,
        "Sau khi xác nhận, số lượng tồn kho sẽ tự động được cập nhật giảm tương ứng tại kho đã chọn.",
        "info")

    doc.add_heading("6.4. Lọc lịch sử giao dịch", level=2)

    add_numbered(doc, 'Nhấn "Bộ lọc" để mở/đóng bộ lọc')
    add_numbered(doc, "Các tiêu chí lọc:")
    add_bullet(doc, "Khoảng thời gian: Chọn ngày bắt đầu → ngày kết thúc", bold_prefix="Khoảng thời gian: ")
    add_bullet(doc, "Loại giao dịch: Tất cả / Nhập kho / Xuất kho", bold_prefix="Loại giao dịch: ")
    add_bullet(doc, "Kho: Lọc theo kho cụ thể (Kho 1 / Kho 2 / Kho 3)", bold_prefix="Kho: ")
    add_numbered(doc, 'Nhấn "Xóa lọc" để bỏ tất cả bộ lọc')

    doc.add_heading("6.5. Xem chi tiết giao dịch", level=2)

    add_body(doc,
        "Nhấn vào bất kỳ giao dịch nào trong danh sách → Mở trang Chi tiết giao dịch. "
        "Hiển thị đầy đủ: loại giao dịch, kho, danh sách sản phẩm, số lượng, giá, tổng cộng, ghi chú, thời gian, người tạo."
    )

    doc.add_heading("6.6. Báo cáo hàng ngày", level=2)

    add_numbered(doc, "Nhấn biểu tượng 📊 (góc trên phải tab Nhật ký)")
    add_numbered(doc, "Chọn ngày: Dùng mũi tên ◀ ▶ để chuyển ngày, hoặc nhấn vào ngày để chọn từ lịch")
    add_numbered(doc, "Xem tổng kết trong ngày:")
    add_bullet(doc, "Thẻ Tổng nhập: Tổng giá trị nhập kho & số phiếu")
    add_bullet(doc, "Thẻ Tổng xuất: Tổng giá trị xuất kho & số phiếu")
    add_bullet(doc, "Chênh lệch: Giá trị Xuất − Nhập (Lãi/Lỗ)")
    add_bullet(doc, "Ghi nợ: Tổng nợ phát sinh trong ngày")
    add_bullet(doc, "Danh sách chi tiết: Từng phiếu nhập/xuất trong ngày")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7: QUẢN LÝ CÔNG NỢ
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("7. Quản lý Công nợ Khách hàng", level=1)

    doc.add_heading("7.1. Xem lịch sử công nợ", level=2)

    add_numbered(doc, "Vào tab Danh mục → Khách hàng")
    add_numbered(doc, 'Nhấn vào khách hàng → Nhấn "Xem lịch sử công nợ"')
    add_numbered(doc, "Trang Công nợ hiển thị:")
    add_bullet(doc, "Tổng công nợ hiện tại (số lớn phía trên)")
    add_bullet(doc, "Danh sách lịch sử: Mỗi bản ghi gồm Loại (Ghi nợ / Thanh toán), Số tiền, Số dư sau giao dịch, Thời gian, Ghi chú")

    doc.add_heading("7.2. Thanh toán công nợ", level=2)

    add_numbered(doc, 'Nhấn nút "💳 Thanh toán"')
    add_numbered(doc, "Nhập Số tiền thanh toán (bắt buộc) và Ghi chú (tùy chọn)")
    add_numbered(doc, 'Nhấn "Xác nhận"')

    doc.add_heading("7.3. Tất toán", level=2)

    add_numbered(doc, 'Nhấn nút "✅ Tất toán"')
    add_numbered(doc, "Xác nhận hành động → Toàn bộ công nợ sẽ được reset về 0")

    add_note_box(doc,
        "Tất toán sẽ xóa toàn bộ công nợ. Đảm bảo đã thu đủ tiền trước khi thực hiện.",
        "danger")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8: ĐỐI SOÁT
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("8. Đối soát Tồn kho", level=1)

    add_body(doc,
        "Chức năng Đối soát giúp bạn kiểm kê hàng hóa thực tế trong kho "
        "và so sánh với dữ liệu hệ thống."
    )

    doc.add_heading("8.1. Cách thực hiện đối soát", level=2)

    add_numbered(doc, 'Từ tab Tồn kho, nhấn nút "✅ Đối soát" (nút nổi phía dưới)')
    add_numbered(doc, "Chọn kho cần kiểm: Tất cả / Kho 1 / Kho 2 / Kho 3")
    add_numbered(doc, "Với mỗi sản phẩm, ứng dụng hiển thị bảng so sánh:")

    add_styled_table(doc,
        ["Cột", "Ý nghĩa"],
        [
            ["Kho", "Tên kho"],
            ["Hệ thống", "Số lượng theo dữ liệu ứng dụng"],
            ["Thực tế", "Ô nhập — Bạn nhập số đếm thực tế"],
            ["Kết quả", "Hiển thị sau khi so sánh (Khớp / Chênh lệch)"],
        ],
        header_color="FBBC04",
        col_widths=[4, 12],
    )

    add_numbered(doc, 'Đi đến từng kho, đếm thực tế và nhập vào ô "Thực tế"')
    add_numbered(doc, 'Nhấn nút "So sánh" ở cuối trang')

    doc.add_heading("8.2. Xem kết quả đối soát", level=2)

    add_body(doc, 'Sau khi nhấn "So sánh", ứng dụng hiển thị:')

    add_bullet(doc, "Bảng tổng kết phía trên:", bold_prefix="Bảng tổng kết: ")
    add_bullet(doc, "  ✅ Khớp: Số vị trí hệ thống = thực tế", level=1)
    add_bullet(doc, "  ❌ Lệch: Số vị trí có chênh lệch", level=1)
    add_bullet(doc, "  🔄 Tổng lệch: Tổng số lượng chênh lệch tuyệt đối", level=1)
    add_bullet(doc, "  📊 Thanh tiến trình: Tỷ lệ phần trăm vị trí đã khớp", level=1)

    add_bullet(doc, "Chi tiết từng sản phẩm:", bold_prefix="Chi tiết: ")
    add_bullet(doc, "  Sản phẩm khớp: Viền xanh, icon ✅", level=1)
    add_bullet(doc, "  Sản phẩm lệch: Viền đỏ, icon ⚠, hiển thị mức chênh lệch (VD: +3 hoặc -2)", level=1)

    add_note_box(doc,
        "Nếu phát hiện chênh lệch, hãy kiểm tra lại hàng hóa hoặc rà soát lịch sử nhập/xuất để tìm nguyên nhân.",
        "tip")

    add_bullet(doc, 'Nhấn "Nhập lại" nếu muốn đối soát lại')
    add_bullet(doc, 'Nhấn "Hoàn thành đối soát" khi đã xong')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9: THỐNG KÊ
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("9. Thống kê & Báo cáo", level=1)

    doc.add_heading("9.1. Truy cập", level=2)
    add_body(doc, "Từ tab Tồn kho → Nhấn biểu tượng 📊 Thống kê (góc trên phải).")

    doc.add_heading("9.2. Chọn kỳ thống kê", level=2)
    add_body(doc, "Nhấn chip tương ứng để chọn kỳ: 7 ngày / 14 ngày / 30 ngày.")

    doc.add_heading("9.3. Nội dung thống kê", level=2)

    doc.add_heading("Thẻ tổng hợp", level=3)
    add_styled_table(doc,
        ["Thẻ", "Ý nghĩa"],
        [
            ["⬇ Tổng nhập", "Tổng giá trị hàng nhập kho trong kỳ"],
            ["⬆ Tổng xuất", "Tổng giá trị hàng xuất kho trong kỳ"],
            ["📈 Lợi nhuận", "Xuất − Nhập (xanh nếu dương, đỏ nếu âm)"],
            ["💰 Ghi nợ", "Tổng công nợ phát sinh trong kỳ"],
        ],
        header_color="34A853",
        col_widths=[4, 12],
    )

    doc.add_heading("Biểu đồ cột — Nhập/Xuất theo ngày", level=3)
    add_bullet(doc, "Cột xanh: Giá trị nhập kho từng ngày")
    add_bullet(doc, "Cột cam: Giá trị xuất kho từng ngày")
    add_bullet(doc, "Nhấn vào cột để xem chi tiết số liệu")

    doc.add_heading("Biểu đồ tròn — Tỷ lệ Nhập/Xuất", level=3)
    add_body(doc, "Thể hiện tỷ lệ phần trăm giữa nhập kho và xuất kho.")

    doc.add_heading("Top 5 sản phẩm giao dịch nhiều nhất", level=3)
    add_body(doc, "Danh sách 5 sản phẩm có tổng giá trị giao dịch cao nhất trong kỳ, xếp hạng từ 1 đến 5.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10: CÀI ĐẶT
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("10. Cài đặt", level=1)

    add_body(doc, "Truy cập bằng cách nhấn biểu tượng ⚙ ở góc trên phải màn hình chính.")

    doc.add_heading("10.1. Tài khoản", level=2)
    add_bullet(doc, "Xem thông tin tài khoản: Tên hiển thị, Email, Trạng thái (Online 🟢)")
    add_bullet(doc, "Nhấn Đăng xuất để thoát tài khoản (có xác nhận)")

    doc.add_heading("10.2. Thông báo", level=2)

    add_styled_table(doc,
        ["Tùy chọn", "Mô tả"],
        [
            ["Nhắc nhở báo cáo hàng ngày", "Bật/Tắt thông báo nhắc nhở tổng kết cuối ngày"],
            ["Giờ nhắc nhở", "Chọn giờ nhận thông báo (mặc định 20:00)"],
            ["Gửi thông báo thử", "Gửi 1 thông báo test để kiểm tra"],
        ],
        col_widths=[5, 11],
    )

    doc.add_heading("10.3. Thông tin ứng dụng", level=2)
    add_bullet(doc, "Phiên bản: 1.0.0", bold_prefix="Phiên bản: ")
    add_bullet(doc, "Mô tả: Quản lý kho hàng pháo hoa — Theo dõi nhập xuất, công nợ, tồn kho", bold_prefix="Mô tả: ")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 11: PHÂN LOẠI NĐ 137
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("11. Phân loại sản phẩm theo Nghị định 137", level=1)

    add_body(doc, "Ứng dụng tuân thủ phân loại theo Nghị định 137/2020/NĐ-CP về quản lý pháo:")

    add_styled_table(doc,
        ["Loại", "Tên gọi", "Quy định"],
        [
            ["A", "Pháo hoa nổ", "🚫 CẤM xuất cho cá nhân"],
            ["B", "Pháo hoa không nổ", "⚠️ Kinh doanh có điều kiện"],
            ["C", "Sản phẩm phụ trợ", "✅ Tự do giao dịch"],
        ],
        header_color="C5221F",
        col_widths=[2, 5, 9],
    )

    add_note_box(doc,
        "Sản phẩm loại A (Pháo hoa nổ) bị cấm xuất cho cá nhân theo quy định pháp luật. "
        "Hệ thống sẽ cảnh báo nếu vi phạm.",
        "danger")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 12: FAQ
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("12. Câu hỏi thường gặp (FAQ)", level=1)

    faqs = [
        (
            "Tại sao tồn kho không cập nhật sau khi nhập/xuất?",
            "Hãy thử kéo xuống (pull-to-refresh) tại tab Tồn kho để tải lại dữ liệu mới nhất."
        ),
        (
            "Tôi có thể quản lý nhiều hơn 3 kho không?",
            "Hiện tại ứng dụng hỗ trợ tối đa 3 kho (Kho 1, Kho 2, Kho 3). Vui lòng liên hệ đội ngũ phát triển nếu cần mở rộng."
        ),
        (
            "Tôi nhập sai giao dịch, làm thế nào để sửa?",
            "Hiện tại ứng dụng chưa hỗ trợ sửa/xóa giao dịch đã tạo. Bạn có thể tạo một giao dịch ngược lại (nhập ↔ xuất) với cùng số lượng để điều chỉnh."
        ),
        (
            "Làm sao biết sản phẩm nào sắp hết hàng?",
            'Tại tab Tồn kho, xem thẻ tóm tắt "Sắp hết" phía trên. Các sản phẩm có số lượng dưới 5 sẽ được đánh dấu đỏ trong danh sách.'
        ),
        (
            "Thông báo nhắc nhở không hoạt động?",
            "1) Kiểm tra trong Cài đặt → Đảm bảo nhắc nhở đang bật. "
            "2) Nhấn 'Gửi thông báo thử' để kiểm tra. "
            "3) Đảm bảo bạn đã cho phép ứng dụng gửi thông báo trong cài đặt hệ thống."
        ),
        (
            "Dữ liệu có được đồng bộ giữa các thiết bị không?",
            "Có! Dữ liệu được lưu trữ trên Firebase Cloud, nên khi đăng nhập cùng tài khoản trên thiết bị khác, dữ liệu sẽ tự động đồng bộ."
        ),
        (
            "Tôi quên mật khẩu, phải làm sao?",
            "Vui lòng liên hệ quản trị viên hệ thống để được hỗ trợ đặt lại mật khẩu."
        ),
    ]

    for question, answer in faqs:
        p = doc.add_paragraph()
        q_run = p.add_run(f"❓ {question}")
        q_run.bold = True
        q_run.font.size = Pt(11)
        q_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        a_p = doc.add_paragraph()
        a_run = a_p.add_run(f"    ➤ {answer}")
        a_run.font.size = Pt(11)
        a_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ══════════════════════════════════════════════════════════════════════════
    # FINAL: LƯU Ý QUAN TRỌNG
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    doc.add_heading("Lưu ý quan trọng", level=1)

    notes = [
        "Luôn thực hiện đối soát tồn kho định kỳ (khuyến nghị hàng tuần) để đảm bảo dữ liệu chính xác.",
        'Kiểm tra thẻ "Sắp hết" hàng ngày để bổ sung hàng kịp thời.',
        "Tuân thủ quy định về phân loại pháo hoa theo Nghị định 137/2020/NĐ-CP.",
        "Đăng xuất khi không sử dụng ứng dụng trên thiết bị công cộng.",
        "Mọi giao dịch đều được ghi nhận người tạo — hãy sử dụng đúng tài khoản cá nhân.",
    ]

    for note in notes:
        add_bullet(doc, f"⚠️  {note}")

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("━" * 40)
    footer_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_p.add_run("Tài liệu hướng dẫn sử dụng — Phiên bản 1.0.0 — 03/04/2026")
    end_run.font.size = Pt(9)
    end_run.font.italic = True
    end_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


# ─── Run ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    output_path = os.path.join(os.path.dirname(__file__), "huong_dan_su_dung.docx")
    doc = create_document()
    doc.save(output_path)
    print(f"[OK] Da tao file thanh cong: {output_path}")
